from pathlib import Path
import sys
sys.path.append(str(Path(__file__).resolve().parent.parent))

import os
from docx import Document
from docxcompose.composer import Composer

def unir_documentos(output_name="manual_final.docx"):
    """
    Une los documentos generados en el siguiente orden:
      1. portada_automatica.docx
      2. FormatoSecciónIntroductoria_automatico.docx
      3. Manual.docx
      4. glosario_automatico.docx
    Conservando formato y estilos con docxcompose.
    """

    uploads_path = Path("uploads")
    templates_path = Path("templates")

    # Orden de documentos
    orden = [
        uploads_path / "portada_automatica.docx",
        uploads_path / "FormatoSecciónIntroductoria_automatico.docx",
        templates_path / "Manual.docx",
        uploads_path / "glosario_automatico.docx",
    ]

    # Verificar existencia
    existentes = [p for p in orden if p.exists()]
    if not existentes:
        raise FileNotFoundError("❌ No se encontró ningún documento para unir.")

    print("📄 Documentos encontrados para unir (en orden):")
    for p in existentes:
        print("   -", p.name)

    # Documento base (la portada)
    base_doc = Document(existentes[0])
    composer = Composer(base_doc)

    # Agregar los demás con salto visual
    for doc_path in existentes[1:]:
        # Añadir salto de página en blanco antes de cada documento
        blank_page = Document()
        blank_page.add_page_break()
        composer.append(blank_page)

        next_doc = Document(doc_path)
        composer.append(next_doc)

    # Guardar resultado final
    final_path = uploads_path / output_name
    composer.save(final_path)
    print(f"✅ Documento final generado: {final_path}")

    # Eliminar temporales excepto Manual.docx y el final
    for path in existentes:
        if path != final_path and path.name != "Manual.docx":
            try:
                os.remove(path)
                print(f"🧹 Eliminado: {path.name}")
            except Exception as e:
                print(f"⚠️ No se pudo eliminar {path.name}: {e}")
        elif path.name == "Manual.docx":
            print(f"🗂 Conservado: {path.name}")

    return final_path


def preparar_y_unir():
    """Ejecuta la unión y limpieza."""
    final_path = unir_documentos()
    print("🧩 Todos los documentos fueron unidos correctamente.")
    return final_path


if __name__ == "__main__":
    preparar_y_unir()

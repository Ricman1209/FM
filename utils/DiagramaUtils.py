from pathlib import Path
import sys
sys.path.append(str(Path(__file__).resolve().parent.parent))

import subprocess
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from prompt.PromptDiagrama import titulo_diagrama, pasos_diagrama


def generar_diagrama_mermaid():
    # 1️⃣ Crear contenido Mermaid dinámico
    contenido = "flowchart TD\n"
    contenido += f'    A[Inicio] --> B["{pasos_diagrama[0]}"]\n'

    for i in range(1, len(pasos_diagrama)):
        letra_anterior = chr(66 + (i - 1))
        letra_actual = chr(66 + i)
        contenido += f'    {letra_anterior}["{pasos_diagrama[i-1]}"] --> {letra_actual}["{pasos_diagrama[i]}"]\n'

    contenido += f'    {chr(66 + len(pasos_diagrama) - 1)} --> Z[Fin]\n'

    # 2️⃣ Definir rutas
    ruta_base = Path("uploads")
    ruta_base.mkdir(exist_ok=True)
    nombre_base = titulo_diagrama.replace(" ", "_")
    ruta_mmd = ruta_base / f"diagrama_{nombre_base}.mmd"
    ruta_png = ruta_base / f"diagrama_{nombre_base}.png"

    # Guardar .mmd
    with open(ruta_mmd, "w", encoding="utf-8") as f:
        f.write(contenido)

    # 3️⃣ Ruta absoluta del ejecutable mmdc.cmd
    ruta_mmdc = r"C:\Users\moral\AppData\Roaming\npm\mmdc.cmd"

    if not Path(ruta_mmdc).exists():
        print("❌ No se encontró el ejecutable mmdc.cmd.")
        print("➡️ Ejecuta: npm install -g @mermaid-js/mermaid-cli")
        return None

    # 4️⃣ Ejecutar Mermaid CLI
    try:
        comando = ["cmd", "/c", ruta_mmdc, "-i", str(ruta_mmd), "-o", str(ruta_png)]
        subprocess.run(comando, check=True)
    except Exception as e:
        print("❌ Error al generar el diagrama con Mermaid CLI:", e)
        return None

    # 5️⃣ Eliminar el archivo .mmd temporal
    try:
        os.remove(ruta_mmd)
    except Exception as e:
        print(f"⚠️ No se pudo eliminar el archivo temporal: {e}")

    return ruta_png


def insertar_diagrama_en_docx(ruta_png):
    """
    Crea un documento Word con la imagen del diagrama ajustada automáticamente.
    """
    if not ruta_png or not Path(ruta_png).exists():
        print("❌ No se encontró la imagen PNG para insertar.")
        return None

    # 📄 Crear documento
    doc = Document()
    doc.add_heading("Diagrama de Flujo", level=1)
    doc.add_paragraph(f"Título del diagrama: {titulo_diagrama}\n")

    # 📐 Ajustar tamaño automáticamente
    imagen = Image.open(ruta_png)
    ancho, alto = imagen.size
    # Ancho máximo permitido dentro del documento (márgenes incluidos)
    ancho_maximo = Inches(6.0)
    proporcion = ancho_maximo / Inches(ancho / 96)  # 96 dpi estimados
    ancho_ajustado = ancho_maximo if proporcion < 1 else Inches(ancho / 96)

    # 📸 Insertar imagen centrada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ruta_png), width=ancho_ajustado)

    # 💾 Guardar resultado
    ruta_docx = Path("uploads") / "diagrama.docx"
    doc.save(ruta_docx)
    print(f"✅ Documento Word generado con el diagrama ajustado: {ruta_docx}")

    return ruta_docx


if __name__ == "__main__":
    ruta_png = generar_diagrama_mermaid()
    if ruta_png:
        insertar_diagrama_en_docx(ruta_png)
